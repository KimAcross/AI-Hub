"""File text extraction utilities for PDF, DOCX, TXT, and MD files."""

from pathlib import Path
from typing import Protocol

import fitz  # PyMuPDF
from docx import Document


class TextExtractor(Protocol):
    """Protocol for text extractors."""

    def extract(self, file_path: Path) -> str:
        """Extract text from a file."""
        ...


class PDFExtractor:
    """Extract text from PDF files using PyMuPDF."""

    def extract(self, file_path: Path) -> str:
        """Extract text from a PDF file.

        Args:
            file_path: Path to the PDF file.

        Returns:
            Extracted text content.
        """
        text_parts: list[str] = []

        with fitz.open(file_path) as doc:
            for page in doc:
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)

        return "\n\n".join(text_parts)


class DOCXExtractor:
    """Extract text from DOCX files using python-docx."""

    def extract(self, file_path: Path) -> str:
        """Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Extracted text content.
        """
        doc = Document(file_path)
        text_parts: list[str] = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n\n".join(text_parts)


class TextFileExtractor:
    """Extract text from TXT and MD files."""

    def extract(self, file_path: Path) -> str:
        """Extract text from a text file.

        Args:
            file_path: Path to the text file.

        Returns:
            File content as text.
        """
        return file_path.read_text(encoding="utf-8")


# Mapping of file extensions to extractors
EXTRACTORS: dict[str, TextExtractor] = {
    "pdf": PDFExtractor(),
    "docx": DOCXExtractor(),
    "txt": TextFileExtractor(),
    "md": TextFileExtractor(),
}

# Allowed file extensions
ALLOWED_EXTENSIONS = set(EXTRACTORS.keys())


def get_file_type(filename: str) -> str | None:
    """Get the file type from a filename.

    Args:
        filename: Name of the file.

    Returns:
        File extension without the dot, or None if not allowed.
    """
    ext = Path(filename).suffix.lower().lstrip(".")
    return ext if ext in ALLOWED_EXTENSIONS else None


def extract_text(file_path: Path, file_type: str) -> str:
    """Extract text from a file based on its type.

    Args:
        file_path: Path to the file.
        file_type: Type of the file (pdf, docx, txt, md).

    Returns:
        Extracted text content.

    Raises:
        ValueError: If file type is not supported.
    """
    extractor = EXTRACTORS.get(file_type)
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")

    return extractor.extract(file_path)
